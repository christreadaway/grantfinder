import io
from typing import Optional
from pypdf import PdfReader
from docx import Document as DocxDocument


class DocumentService:
    """Service for extracting text from uploaded documents."""

    MAX_PAGES = 50  # Limit for large documents

    @staticmethod
    async def extract_text(file_content: bytes, file_type: str, filename: str) -> tuple[str, str]:
        """
        Extract text from a document.
        Returns (extracted_text, document_type_description)
        """
        if file_type == "pdf":
            return await DocumentService._extract_from_pdf(file_content, filename)
        elif file_type == "docx":
            return await DocumentService._extract_from_docx(file_content, filename)
        elif file_type == "txt":
            return await DocumentService._extract_from_txt(file_content, filename)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    async def _extract_from_pdf(file_content: bytes, filename: str) -> tuple[str, str]:
        """Extract text from PDF file."""
        try:
            reader = PdfReader(io.BytesIO(file_content))
            text_parts = []

            # Limit to MAX_PAGES
            num_pages = min(len(reader.pages), DocumentService.MAX_PAGES)

            for i in range(num_pages):
                page = reader.pages[i]
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            extracted_text = "\n\n".join(text_parts)

            # Determine document type from filename
            doc_type = DocumentService._guess_document_type(filename)

            return extracted_text, doc_type

        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    async def _extract_from_docx(file_content: bytes, filename: str) -> tuple[str, str]:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            extracted_text = "\n\n".join(text_parts)
            doc_type = DocumentService._guess_document_type(filename)

            return extracted_text, doc_type

        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    async def _extract_from_txt(file_content: bytes, filename: str) -> tuple[str, str]:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    extracted_text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                extracted_text = file_content.decode('utf-8', errors='ignore')

            doc_type = DocumentService._guess_document_type(filename)
            return extracted_text, doc_type

        except Exception as e:
            raise ValueError(f"Failed to extract text from TXT: {str(e)}")

    @staticmethod
    def _guess_document_type(filename: str) -> str:
        """Guess the document type from filename."""
        filename_lower = filename.lower()

        if "bulletin" in filename_lower:
            return "weekly bulletin"
        elif "minute" in filename_lower:
            return "meeting minutes"
        elif "newsletter" in filename_lower:
            return "newsletter"
        elif "strategic" in filename_lower or "plan" in filename_lower:
            return "strategic plan"
        elif "budget" in filename_lower:
            return "budget document"
        elif "annual" in filename_lower and "report" in filename_lower:
            return "annual report"
        elif "capital" in filename_lower or "campaign" in filename_lower:
            return "capital campaign materials"
        elif "curriculum" in filename_lower:
            return "curriculum plan"
        elif "enrollment" in filename_lower:
            return "enrollment report"
        else:
            return "parish document"

    @staticmethod
    def get_file_type(filename: str) -> Optional[str]:
        """Get file type from filename."""
        filename_lower = filename.lower()
        if filename_lower.endswith('.pdf'):
            return 'pdf'
        elif filename_lower.endswith('.docx'):
            return 'docx'
        elif filename_lower.endswith('.txt'):
            return 'txt'
        return None
