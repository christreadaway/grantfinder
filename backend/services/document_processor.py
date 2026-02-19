"""
Document processor for GrantFinder AI.
Handles PDF, DOCX, and TXT file text extraction.
"""
import io
import logging
from typing import Optional

from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


async def process_document(content: bytes, file_type: str) -> str:
    """
    Extract text from uploaded document.

    Args:
        content: Raw file bytes
        file_type: File extension (.pdf, .docx, .txt)

    Returns:
        Extracted text content
    """
    try:
        if file_type == '.pdf':
            return extract_pdf_text(content)
        elif file_type == '.docx':
            return extract_docx_text(content)
        elif file_type == '.txt':
            return extract_txt_text(content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    except Exception as e:
        logger.error(f"Document processing error: {e}")
        raise


def extract_pdf_text(content: bytes) -> str:
    """Extract text from PDF file."""
    try:
        reader = PdfReader(io.BytesIO(content))
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from PDF ({len(reader.pages)} pages)")
        return full_text

    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_docx_text(content: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        doc = Document(io.BytesIO(content))
        text_parts = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from DOCX")
        return full_text

    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_txt_text(content: bytes) -> str:
    """Extract text from TXT file."""
    try:
        # Try common encodings
        for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
            try:
                text = content.decode(encoding)
                logger.info(f"Extracted {len(text)} characters from TXT ({encoding})")
                return text
            except UnicodeDecodeError:
                continue

        raise ValueError("Could not decode text file with any common encoding")

    except Exception as e:
        logger.error(f"TXT extraction error: {e}")
        raise ValueError(f"Failed to extract text from TXT: {str(e)}")
