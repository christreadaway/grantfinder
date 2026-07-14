"""
GrantWriter export builders (PRD 4.8).

Formats: Word (.docx), Markdown, plain text with per-field word/char counts
(for portal paste-in), and a form-field mapping. PDF is intentionally not
generated here (no headless browser in this stack); Word/Markdown cover the
attach-a-document cases and are trivially printable to PDF.

Hard rule enforced by the caller: nothing over a section's word/char limit is
ever exported.
"""
import io
import json
from datetime import datetime
from typing import List, Tuple

from models.writer_schemas import Application, GrantSpec, SectionDraft


def _ordered(drafts: List[SectionDraft], spec: GrantSpec) -> List[SectionDraft]:
    """Order drafts to match the grant spec's section order."""
    order = {s.id: i for i, s in enumerate(spec.required_sections)}
    return sorted(drafts, key=lambda d: order.get(d.section_id, 999))


def _submission_checklist(spec: GrantSpec) -> List[str]:
    items = [f"Format: {c}" for c in spec.format_constraints]
    items += [f"Deliverable: {d}" for d in spec.deliverables]
    return items


def build_markdown(app: Application, spec: GrantSpec, drafts: List[SectionDraft]) -> str:
    lines = [
        f"# {app.grant_name}",
        f"**Funder:** {app.funder}  ",
        f"**Deadline:** {app.deadline or 'Check guidelines'}  ",
        f"**Exported:** {datetime.utcnow().strftime('%Y-%m-%d')}",
        "",
    ]
    for d in _ordered(drafts, spec):
        lines.append(f"## {d.title}")
        lines.append("")
        lines.append(d.current_draft)
        lines.append("")
    checklist = _submission_checklist(spec)
    if checklist:
        lines.append("---")
        lines.append("## Submission Checklist")
        lines += [f"- [ ] {item}" for item in checklist]
    return "\n".join(lines)


def build_plain_text(app: Application, spec: GrantSpec, drafts: List[SectionDraft]) -> str:
    """Portal paste-in format: each field labeled with its counts."""
    blocks = [
        f"{app.grant_name} - {app.funder}",
        f"Deadline: {app.deadline or 'Check guidelines'}",
        "=" * 60,
    ]
    limits = {s.id: s for s in spec.required_sections}
    for d in _ordered(drafts, spec):
        section = limits.get(d.section_id)
        limit_bits = []
        if section and section.word_limit:
            limit_bits.append(f"limit {section.word_limit} words")
        if section and section.char_limit:
            limit_bits.append(f"limit {section.char_limit} chars")
        limit_note = f" ({'; '.join(limit_bits)})" if limit_bits else ""
        blocks.append("")
        blocks.append(f"FIELD: {d.title}{limit_note}")
        blocks.append(f"[{d.word_count} words / {d.char_count} characters]")
        blocks.append("-" * 60)
        blocks.append(d.current_draft)
    return "\n".join(blocks)


def build_form_map(app: Application, spec: GrantSpec, drafts: List[SectionDraft]) -> str:
    """Field-by-field JSON map for form-based portals (transcribe now, auto-populate later)."""
    limits = {s.id: s for s in spec.required_sections}
    fields = []
    for d in _ordered(drafts, spec):
        section = limits.get(d.section_id)
        fields.append({
            "field": d.title,
            "prompt": section.prompt if section else "",
            "word_limit": section.word_limit if section else None,
            "char_limit": section.char_limit if section else None,
            "word_count": d.word_count,
            "char_count": d.char_count,
            "content": d.current_draft,
        })
    return json.dumps({
        "grant": app.grant_name,
        "funder": app.funder,
        "deadline": app.deadline,
        "exported_at": datetime.utcnow().isoformat(),
        "fields": fields,
        "submission_checklist": _submission_checklist(spec),
    }, indent=2)


def build_docx(app: Application, spec: GrantSpec, drafts: List[SectionDraft]) -> bytes:
    """Clean, formatted Word document ready to attach."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(app.grant_name, level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Funder: {app.funder}    Deadline: {app.deadline or 'Check guidelines'}").italic = True

    for d in _ordered(drafts, spec):
        doc.add_heading(d.title, level=1)
        for para in d.current_draft.split("\n\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                for run in p.runs:
                    run.font.size = Pt(11)

    checklist = _submission_checklist(spec)
    if checklist:
        doc.add_page_break()
        doc.add_heading("Submission Checklist", level=1)
        for item in checklist:
            doc.add_paragraph(item, style="List Bullet")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_export(
    fmt: str,
    app: Application,
    spec: GrantSpec,
    drafts: List[SectionDraft],
) -> Tuple[bytes, str, str]:
    """Returns (content_bytes, filename, media_type)."""
    safe_name = "".join(c if c.isalnum() or c in "-_ " else "" for c in app.grant_name)[:50].strip().replace(" ", "_")
    stamp = datetime.utcnow().strftime("%Y%m%d")

    if fmt == "docx":
        return (
            build_docx(app, spec, drafts),
            f"{safe_name}_application_{stamp}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if fmt == "md":
        return (
            build_markdown(app, spec, drafts).encode("utf-8"),
            f"{safe_name}_application_{stamp}.md",
            "text/markdown",
        )
    if fmt == "txt":
        return (
            build_plain_text(app, spec, drafts).encode("utf-8"),
            f"{safe_name}_application_{stamp}.txt",
            "text/plain",
        )
    if fmt == "form_map":
        return (
            build_form_map(app, spec, drafts).encode("utf-8"),
            f"{safe_name}_form_map_{stamp}.json",
            "application/json",
        )
    raise ValueError(f"Unsupported export format: {fmt}")
